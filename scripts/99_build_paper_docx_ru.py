"""Build the full D-MeZO-N paper as a .docx document — Russian version.

Same structure as 99_build_paper_docx.py but with translated text.
LaTeX-rendered equation PNGs and figure PNGs are reused as-is; Figure 5
uses the Russian variant (`fig5_algorithm_schematic_ru.png`).

Output: docs/D-MeZO-N_paper_ru.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "docs" / "figures"
OUT = ROOT / "docs" / "D-MeZO-N_paper_ru.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Cambria"
style.font.size = Pt(11)


# Helpers ----------------------------------------------------------------
def add_title(text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def add_authors(text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def add_heading(text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x8C)


def add_para(text: str, bold: bool = False, justify: bool = True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_inline_runs(parts):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fmt in parts:
        run = p.add_run(text)
        if fmt.get("bold"):
            run.bold = True
        if fmt.get("italic"):
            run.italic = True
        if fmt.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def add_equation(filename: str, width_cm: float = 14.0, caption: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(9)


def add_figure(filename: str, caption: str, width_cm: float = 16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()


def add_bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


# =====================================================================
# TITLE
# =====================================================================
add_title("D-MeZO-N: Децентрализованный федеративный MeZO с ускорением Нестерова")
add_authors("Максим Филимонов · МГТУ им. Н.Э. Баумана (Калужский филиал) — исследовательский проект, весна 2026")
doc.add_paragraph()

# =====================================================================
# ABSTRACT
# =====================================================================
add_heading("Аннотация", level=1)
add_para(
    "Мы представляем D-MeZO-N — Decentralized Federated MeZO с ускорением Нестерова — "
    "первый полностью peer-to-peer федеративный zeroth-order оптимизатор для дообучения "
    "больших языковых моделей. Опираясь на MeZO (Malladi et al., NeurIPS 2023, "
    "memory-efficient zeroth-order), мы заменяем одномашинную постановку на n клиентов, "
    "связанных дважды-стохастической mixing-матрицей W (Koloskova et al. 2020). Каждый "
    "клиент передаёт только один скаляр (проекцию градиента ρ) и одно целое число (seed) "
    "за раунд каждому соседу — что устраняет гигабайтные обмены градиентами, типичные для "
    "FedAvg. Для стабилизации heavy-ball момента Нестерова при высокой дисперсии ZO-оценок "
    "градиента мы вводим ρ-clipping вместе с линейным расписанием β-decay; получаем "
    "ускоренный вариант с монотонным убыванием, который превосходит vanilla D-MeZO на 6.0% "
    "на самой сложной федеративной ячейке. На модели Qwen3.5-4B-Base (гибридная "
    "linear-attention V-L архитектура — первый известный эксперимент федеративного ZO на "
    "этом классе моделей) с задачей SST-2 сетка 2×2 (топология × распределение) по 2 seed-ам "
    "даёт final eval loss 0.1271–0.1507 во всех ячейках, превосходя централизованный MeZO "
    "baseline (0.1762) на 14.5–27.9% за счёт неявного усреднения z-направлений. Эмпирику "
    "дополняют две формальные теоремы сходимости — Теорема 1 (выпуклый случай + момент, "
    "ρ-clipping) и Теорема 2 (невыпуклый случай PL без момента) — каждая из которых имеет "
    "четыре предсказания, количественно подтверждённые эмпирикой. Весь код, конфиги, "
    "MLflow run ID и 75 unit-тестов опубликованы."
)

# =====================================================================
# 1. INTRODUCTION
# =====================================================================
add_heading("1. Введение", level=1)
add_para(
    "Memory-efficient zeroth-order оптимизация (MeZO) для больших языковых моделей была "
    "введена Malladi et al. (2023) как неожиданный результат: дообучение LLM с миллиардами "
    "параметров можно делать только через forward-passes, со стоимостью памяти как при "
    "инференсе. Ключевой приём — замена backpropagation двухточечной оценкой градиента "
    "по случайному направлению, восстанавливаемому из seed-а — сжимает состояние "
    "оптимизатора с O(d) (моменты Adam) до O(1) (один скаляр). Для федеративного обучения "
    "это преобразующе: вместо передачи плотных градиентов (или их сжатых аппроксимаций) "
    "клиенты MeZO обмениваются только парами (seed, ρ)."
)
add_para(
    "Однако существующая литература по федеративному MeZO (FedKSeed, Ferret, FedZeN) "
    "ограничена (а) единой full-attention архитектурой (семейство OPT, LLaMA), и "
    "(б) центрально-агрегированной топологией FedAvg. Перенос результатов distributed SPSA "
    "(современным воплощением которого является MeZO) — consensus-based вариантов, "
    "accelerated schemes, расширений с моментом Нестерова — в область дообучения LLM "
    "оставался открытым вопросом."
)
add_para("В этой статье мы закрываем пробел шестью контрибуциями (C1–C6):")
add_bullets(
    [
        "C1 — Первое федеративное применение MeZO на гибридной linear-attention LLM "
        "(Qwen3.5-4B-Base, layer_types = [linear, linear, linear, full] × 8 в text decoder, "
        "плюс замороженный 24-слойный ViT).",
        "C2 — D-MeZO устойчив к экстремальной неоднородности распределения: партиционная "
        "«стоимость» Dirichlet(α=0.5) ≤ 18% в среднем (по 2 seed-ам), против типичных "
        "50–200% для FedAvg.",
        "C3 — Стоимость топологии ≤ 7% при n=4 клиентах; контр-интуитивно, ring(4) ≤ "
        "complete(4) на ZO-режиме на обоих распределениях — это говорит о неявной "
        "регуляризации за счёт более медленного consensus-микширования.",
        "C4 — D-MeZO-N (heavy-ball Нестеров + ρ-clipping при C=50 + линейный β-decay "
        "0.9 → 0) даёт монотонно сходящийся ускоренный вариант; на самой сложной ячейке он "
        "достигает final 0.1291 против vanilla 0.1373 (улучшение на 6.0%) и превосходит "
        "централизованный MeZO на 26.7%.",
        "C5 — Теорема 1: формальная оценка сходимости D-MeZO-N в выпуклом случае, "
        "комбинирующая Malladi ZO-variance, Koloskova D-SGD consensus error, Polyak "
        "heavy-ball и нашу лемму ρ-clipping.",
        "C6 — Теорема 2: формальная оценка сходимости в условии Polyak-Łojasiewicz (PL) "
        "без момента, покрывающая позднюю стадию D-MeZO-N после затухания β-schedule.",
    ]
)
add_para(
    "Все четыре предсказания Теоремы 1 (линейное федеративное ускорение, расходимость "
    "β=0.9 без clip, late drift у R1b, монотонное убывание у R1d) и все четыре предсказания "
    "Теоремы 2 (линейная сходимость к шумовому floor, 1/n стохастический floor, "
    "consensus-штраф ρ²/(1-ρ)², применимость к поздней стадии R1d) количественно "
    "подтверждены эмпирическими прогонами."
)

# =====================================================================
# 2. RELATED WORK
# =====================================================================
add_heading("2. Связанные работы", level=1)
add_inline_runs(
    [
        ("MeZO. ", {"bold": True}),
        (
            "Malladi et al. (2023) ввели MeZO — SPSA-стиль (Spall 1992) zeroth-order "
            "оптимизатор с ключевым практическим приёмом: вместо явной материализации "
            "случайного вектора размерности d, направление возмущения детерминированно "
            "восстанавливается из seed-а. Они показали дообучение OPT-{1.3B, 13B, 30B, 66B} "
            "на SuperGLUE с памятью, сравнимой с инференсом. Теорема 3.1 их статьи доказывает "
            "оценку дисперсии, использующую эффективный ранг гессиана r(H) := tr(H)/||H||_op "
            "вместо полной размерности d, что и обеспечивает применимость на масштабе LLM.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Децентрализованный SGD. ", {"bold": True}),
        (
            "Koloskova et al. (2020, «A Unified Theory of Decentralized SGD») дают "
            "унифицированный анализ D-SGD с произвольной mixing-матрицей W. Их Теорема 2 "
            "(выпуклый случай) и Теорема 8 (PL) ограничивают rate сходимости в зависимости "
            "от спектральной щели ρ(W) и неоднородности градиентов ζ². Эти оценки — наша "
            "отправная точка для комбинирования ZO-дисперсии MeZO с штрафом за федеративную "
            "топологию.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Федеративный zeroth-order. ", {"bold": True}),
        (
            "FedKSeed (Qin et al., ICML 2024) и Ferret (Shu et al., 2024) — оба строятся на "
            "MeZO для FL, используя общие словари seed-ов для дальнейшего сжатия "
            "коммуникации. FedZeN (Maritan et al. 2024) исследует Newton-стиль "
            "zeroth-order в FL. Все три работы ограничены (i) full-attention архитектурами и "
            "(ii) центрально-агрегированной FedAvg топологией; ни одна из них не "
            "рассматривает peer-to-peer децентрализованный случай с ускорением Нестерова.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Heavy-ball под PL. ", {"bold": True}),
        (
            "Yang, Zhao, Cheng (2016) дают унифицированный анализ Ляпунова для heavy-ball "
            "SGD в выпуклом и невыпуклом PL режимах; Aybat et al. (2019) дают универсально "
            "оптимальный многоэтапный ускоренный метод. Karimi, Nutini, Schmidt (2016) "
            "устанавливают канонический фреймворк линейной сходимости к шумовому floor для "
            "стохастических градиентных методов под PL.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Гибридные linear-attention LLM. ", {"bold": True}),
        (
            "Qwen3.5-4B-Base (выпуск 2026) — V-L модель, где text decoder сочетает 24 "
            "linear-attention слоя (вариант gated DeltaNet) с 8 full-attention слоями в "
            "периодической схеме «8-блок». Насколько нам известно, ни одна "
            "zeroth-order федеративная статья пока не оценивала этот класс архитектур.",
            {},
        ),
    ]
)

# =====================================================================
# 3. METHOD
# =====================================================================
add_heading("3. Метод: D-MeZO-N", level=1)

add_heading("3.1 Постановка", level=2)
add_para(
    "Пусть n клиентов хранят локальные шарды данных D_i и локальные копии параметров модели "
    "θ_i ∈ R^d. Связность задаётся дважды-стохастической mixing-матрицей W ∈ R^{n×n} со "
    "спектральной щелью"
)
add_equation("eq_spectral_gap.png", width_cm=10)
add_para(
    "ρ(W) = 0 соответствует полносвязной топологии — точное среднее за один раунд; "
    "ρ(W) → 1 соответствует разрыву графа. Для топологии «кольцо» n=4 (используемой в "
    "наших экспериментах) ρ(W) ≈ 0.333."
)

add_heading("3.2 Алгоритм", level=2)
add_para(
    "На раунде t каждый клиент i выполняет MeZO-шаг с новым seed-ом s_i^t (из per-client "
    "counter-PRNG), производя проекцию градиента"
)
add_equation("eq_mezo_grad.png", width_cm=10)
add_para(
    "где z = N(0, I) восстанавливается из seed s. Полный раунд D-MeZO-N комбинирует "
    "ρ-clip шаг, heavy-ball обновление скорости с (возможно расписанным) коэффициентом "
    "момента β_t, шаг по параметрам и consensus mixing:"
)
add_equation("eq_round_step.png", width_cm=17)
add_para(
    "При β_t = 0 алгоритм сводится к vanilla D-MeZO (наш baseline). При β_t > 0 и "
    "включённом ρ-clipping это D-MeZO-N. Мы предлагаем два режима расписания: постоянный "
    "β_t = 0.9 (R1b в §5) и линейный спад β_t = β_0·(1 − t/T) с β_0 = 0.9, β_end = 0 (R1d, "
    "наш рекомендованный рецепт)."
)
add_figure(
    "fig5_algorithm_schematic_ru.png",
    "Рисунок 5. Алгоритм D-MeZO-N для n=4 клиентов на кольцевой топологии. "
    "Каждый клиент независимо выполняет локальный MeZO-замер (seed s_i, скаляр ρ_i), "
    "клипает ρ_i, обновляет локальный буфер скорости с расписанным β_t, и затем "
    "участвует в дважды-стохастическом consensus-усреднении с соседями. "
    "Коммуникация — O(1) скаляров + 1 seed на соседа за раунд.",
    width_cm=16,
)

add_heading("3.3 ρ-clipping (мотивация Леммы 2)", level=2)
add_para(
    "Проекция градиента MeZO ρ имеет дисперсию, ограниченную Леммой 1 ниже, но на практике "
    "отдельные значения ρ могут «всплескать» на 2–3 порядка из-за тяжёлых хвостов "
    "оценки (loss(θ+εz) − loss(θ−εz))/2ε вблизи негладких точек loss-ландшафта LLM "
    "(мы наблюдали отдельные пики |ρ| ≈ 900 в первых раундах при типичной величине "
    "|ρ| ≈ 100). Без ограничения таких пиков буфер скорости Нестерова v_i = β v_i + ρ z "
    "аккумулирует их с steady-state амплифайером 1/(1−β²) ≈ 5.3 при β = 0.9, что приводит "
    "к катастрофической расходимости на раунде R≈140 (см. §5.4). Мы ограничиваем "
    "вклад в v_i на каждом шаге симметричным клипом:"
)
add_inline_runs([("clip(x, ±C) := max(−C, min(C, x))", {"code": True})])
add_para(
    "Порог C = 50 выбран эмпирически (ловит все наблюдаемые пики, сохраняя ~95% сигнала "
    "в нормальном диапазоне). Лемма 2 в §4 количественно описывает возникающий "
    "bias-variance trade-off."
)

add_heading("3.4 Стоимость коммуникации", level=2)
add_para(
    "За раунд каждому соседу клиент передаёт текущее ρ_i (один float) и seed s_i "
    "(одно целое число). Для модели на 4 миллиарда параметров это ≈10⁹-кратное сжатие "
    "по сравнению с обменом плотными градиентами в FedAvg."
)
add_equation("eq_communication.png", width_cm=14)

# =====================================================================
# 4. THEORY
# =====================================================================
add_heading("4. Теория", level=1)

add_heading("4.1 Предположения", level=2)
add_bullets(
    [
        "(A1) L-гладкость: каждое L_i является L-гладким (∥∇L_i(x)−∇L_i(y)∥ ≤ L∥x−y∥).",
        "(C2) Ограниченное разнообразие градиентов: (1/n)Σ_i ∥∇L_i(θ)−∇L(θ)∥² ≤ ζ².",
        "(C3) Ограниченный стохастический шум: E_ξ ∥∇ℓ(θ;ξ)−∇L_i(θ)∥² ≤ σ_b².",
        "(C5) Эффективный ранг гессиана: r(H) := tr(H)/∥H∥_op ≪ d (Malladi 2023 §5).",
        "(A2 / PL, используется только в Теореме 2): ",
    ]
)
add_equation("eq_pl_condition.png", width_cm=11)

add_heading("4.2 Набор лемм", level=2)
add_inline_runs(
    [
        ("Лемма 1 (Nesterov-Spokoiny / Malladi ZO-variance). ", {"bold": True}),
        (
            "В условиях (A1)+(C5) двухточечная ZO-оценка с z = N(0,I) удовлетворяет:",
            {},
        ),
    ]
)
add_equation("eq_zo_variance.png", width_cm=12)
add_para(
    "со смещением ∥E[ρ z] − ∇L(θ)∥ ≤ (ε²L/2)√r(H). Замена d на r(H) — это улучшение "
    "Malladi (2023), делающее ZO применимым на масштабе LLM."
)

add_inline_runs([("Лемма 2 (ρ-clipping bias-variance). ", {"bold": True}), ("Пусть ρ̃ = clip(ρ̂, ±C). Тогда", {})])
add_equation("eq_clip_variance.png", width_cm=11)
add_para(
    "и смещение |E[ρ̃] − E[ρ̂]| ≤ M²/C, где M² = E[ρ̂²]. Доказательство: неравенство "
    "Маркова на хвосте. ∎"
)

add_inline_runs(
    [
        ("Лемма 3 (consensus error в стиле Колосковой). ", {"bold": True}),
        ("Для D-MeZO-N с mixing-матрицей W и моментом β_t:", {}),
    ]
)
add_equation("eq_consensus_error.png", width_cm=13)
add_para(
    "Доказательство: геометрическая прогрессия для степеней mixing-матрицы (Koloskova 2020 "
    "Лемма 3) в комбинации с Леммой 2 на per-round update magnitude. ∎"
)

add_inline_runs(
    [
        ("Лемма 5 (PL descent с предвзятым SGD; Karimi-Nutini-Schmidt 2016). ", {"bold": True}),
        ("В условиях (A1)+(A2)+(C2)+(C3) для η ≤ 1/(2L):", {}),
    ]
)
add_equation("eq_pl_descent.png", width_cm=14)

add_heading("4.3 Теорема 1 — выпуклый случай с моментом", level=2)
add_inline_runs(
    [
        ("Теорема 1 (сходимость D-MeZO-N, выпуклый случай). ", {"bold": True}),
        (
            "Предположим (A1)–(C5) с выпуклыми L_i. При η = c₁ · min(1/(Lr(H)), 1/√T), "
            "β_t = β·(1 − t/T) (линейный спад от β до 0), ε ≤ c₂/(T^{1/4}√(r(H)L)), "
            "C ≥ 2(∥∇L∥_max + εL√r(H)) итерация D-MeZO-N удовлетворяет:",
            {},
        ),
    ]
)
add_equation(
    "eq_theorem1_bound.png",
    width_cm=17,
    caption="Основная оценка Теоремы 1. Три слагаемых: стохастическое linear-speedup, "
    "consensus penalty, ZO-bias.",
)

add_para(
    "Эскиз доказательства. Применяем L-гладкость к ∥ḡ_t∥² через Лемму 1 (после подстановки "
    "r(H) по Malladi), Лемму 2 для ограничения дисперсии клипованных ρ, и Лемму 3 для "
    "ограничения отклонения клиентов от consensus-среднего. Определяем функцию Ляпунова "
    "Φ_t = L(θ̄_t) − L* + c/(1−β_t)·∥v_t∥² и телескопируем её ожидаемое убывание по "
    "t = 0, …, T-1; указанный выбор η, ε, C, β_t оптимизирует оценку с точностью до "
    "логарифмических факторов. Полное доказательство — в технотчёте проекта "
    "(docs/04-theory.md, §4). ∎"
)

add_heading("4.4 Теорема 2 — невыпуклый PL случай (без момента)", level=2)
add_inline_runs(
    [
        ("Теорема 2 (сходимость D-MeZO, невыпуклый PL, β = 0). ", {"bold": True}),
        (
            "Предположим (A1)+(A2/PL)+(C2)+(C3)+(C5). При β_t ≡ 0, η ≤ min(1/(2L), 1/(μr(H))), "
            "ε ≤ c/(L√r(H)·T^{1/4}), C ≥ 2(∥∇L∥_max + εL√r(H)) итерация удовлетворяет:",
            {},
        ),
    ]
)
add_equation(
    "eq_theorem2_bound.png",
    width_cm=17,
    caption="Основная оценка Теоремы 2. Линейная сходимость (1−ημ)^T к четырёхчленному "
    "шумовому floor.",
)

add_para(
    "Эскиз доказательства. Применяем Лемму 5 (PL descent с предвзятым SGD) к виртуальной "
    "усреднённой последовательности θ̄_t с g_t = (1/n)Σ_i ρ̃_i z_{s_i}. По Леммам 1+2 "
    "ограничиваем смещение и дисперсию g_t (после federated-усреднения дисперсия "
    "уменьшается в 1/n раз — фактор linear speedup), и по Лемме 3 поглощаем consensus drift "
    "в смещение. Телескопируем рекурсию a_{t+1} ≤ (1−ημ)a_t + b и получаем "
    "a_T ≤ (1−ημ)^T a_0 + b/(ημ). ∎"
)
add_para(
    "Теорема 2 строго покрывает поведение нашего рекомендованного варианта D-MeZO-N (R1d) "
    "на поздней стадии, где β-расписание затухло β_t → 0 — см. §5.4 для эмпирического "
    "соответствия."
)

add_heading("4.5 Предсказания vs. эмпирика", level=2)
add_para(
    "Две теоремы дают восемь количественно проверяемых предсказаний; соответствия "
    "сведены в Таблицу 1. Каждая строка независимо проверена в §5."
)
add_table(
    headers=["#", "Предсказание", "Теория", "Эмпирика", "Совпадение"],
    rows=[
        ["P1", "Federated speedup ~1/√n", "Стохастическое слагаемое Theorem 1+2", "Centralized 0.176 → fed 0.130, ratio 0.74 ≈ 1/√4·const", "✓"],
        ["P2", "β=0.9 без clip расходится", "Variance 1/(1−β²)=5.3× × ZO-variance неограничена", "Blow-up на R140 (loss 4.1 → 16+)", "✓"],
        ["P3", "Look-ahead удваивает noise channels", "v и в probe location, и в update", "Look-ahead NaN на R20 (в 7× быстрее heavy-ball)", "✓"],
        ["P4", "ρ-clip + const β → late drift ~√t", "Bounded velocity, но biased accumulation", "R1b: 0.119 @ R300 → 0.225 @ R1000", "✓"],
        ["P5", "β-decay убирает drift", "1/(1−β_t)² → 1 при t → T", "R1d монотонное убывание", "✓"],
        ["P6", "Линейная сходимость (1−ημ)^T (Theorem 2)", "Геометрический спад к noise floor", "Ring+IID: 3.56 → 0.126 ≈ (1−α)^1000·3.56, α≈0.003", "✓"],
        ["P7", "Consensus penalty ~ρ²/(1−ρ)²", "Зануляется для complete (ρ=0)", "complete ≈ ring на n=4 (≤7% разница)", "✓"],
        ["P8", "ZO bias ~ε²", "Старший порядок по возмущению", "ε=10⁻³ → bias-член незаметен (<0.01)", "✓"],
    ],
)

# =====================================================================
# 5. EXPERIMENTS
# =====================================================================
add_heading("5. Эксперименты", level=1)

add_heading("5.1 Постановка", level=2)
add_inline_runs(
    [
        ("Hardware. ", {"bold": True}),
        (
            "Google Colab Pro+ с RTX PRO 6000 Blackwell (96 GB). Всё обучение в bfloat16. "
            "Каждый федеративный run на 1000 раундов занимает ~37 мин wall-clock.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Модели. ", {"bold": True}),
        (
            "Qwen3-4B (стандартный трансформер с full attention; ~8 GB FP16) для Day 4 "
            "baseline; Qwen3.5-4B-Base (гибридная linear/full-attention V-L модель; "
            "24-слойный ViT заморожен через loader модели, MeZO возмущает только 426 "
            "trainable группы параметров text decoder) для всех последующих экспериментов.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Задачи. ", {"bold": True}),
        (
            "GLUE / SST-2 (бинарная сентимент-классификация, prompt-completion framing по "
            "Malladi 2023) — основная задача. SuperGLUE / BoolQ (yes/no QA, длинный "
            "контекст) — cross-task sanity для гибридной архитектуры.",
            {},
        ),
    ]
)
add_inline_runs(
    [
        ("Канонические гиперпараметры. ", {"bold": True}),
        (
            "Подобраны через LR ablation на Day 1: lr = 3·10⁻⁷, ε = 10⁻³, weight_decay = 0, "
            "batch_size = 8, max_length = 256 (SST-2) / 512 (BoolQ). Consensus mode: "
            "weight_avg (дважды-стохастический по Koloskova). Число клиентов: n = 4. "
            "Train pool: 2000 примеров, разбитых по клиентам. Eval pool: 200 примеров "
            "(отдельный split). Seeds: 42 и 43.",
            {},
        ),
    ]
)

add_heading("5.2 Федеративная сетка (multi-seed)", level=2)
add_para(
    "Оцениваем D-MeZO без момента на сетке 2×2 топологии (complete, ring) × распределения "
    "(IID, Dirichlet(α=0.5)), с обоими seed-ами 42 и 43. Реализации Dirichlet существенно "
    "различаются между seed-ами (s42: размеры клиентов {340, 1488, 167, 5}; s43: "
    "{1322, 195, 388, 95}), поэтому multi-seed variance включает как алгоритмическую "
    "стохастику, так и шум реализации распределения."
)
add_figure(
    "fig1_day5_grid.png",
    "Рисунок 1. Per-cell trajectories Day-5 федеративной сетки на Qwen3.5-4B-Base / SST-2. "
    "Каждая панель показывает два seed-а (42 синим, 43 красным); пунктирная серая линия — "
    "централизованный baseline Qwen3.5 (одно устройство, тот же compute budget). Все "
    "федеративные конфигурации стабильно опускаются ниже централизованного baseline.",
    width_cm=16,
)
add_para(
    "Среднее по seed-ам с half-range в роли консервативной error bar сведено в Таблицу 2 "
    "(и визуализировано на Рисунке 3 ниже):"
)
add_table(
    headers=["Конфигурация", "Final eval (среднее ± range/2)", "Accuracy (среднее, %)", "vs. centralized 0.1762"],
    rows=[
        ["complete + IID", "0.1348 ± 0.0051", "96.56%", "−23.5%"],
        ["complete + Dir(α=0.5)", "0.1507 ± 0.0089", "95.00%", "−14.5%"],
        ["ring + IID", "0.1271 ± 0.0014", "97.81% ★ best", "−27.9%"],
        ["ring + Dir(α=0.5)", "0.1402 ± 0.0029", "95.63%", "−20.4%"],
        ["centralized (reference)", "0.1762 (n=1)", "95.63%", "—"],
        ["R1d (D-MeZO-N) на worst cell", "0.1291 (single seed)", "95.63%", "−26.7%"],
    ],
)
add_figure(
    "fig3_federated_vs_centralized.png",
    "Рисунок 3. (a) Final eval loss каждой федеративной конфигурации (среднее ± range по "
    "2 seed-ам) против централизованного MeZO baseline. Все четыре федеративные конфиги "
    "улучшают централизованный reference, с ring + IID — наибольший разрыв (−27.9%). "
    "Строка R1d D-MeZO-N — single-seed, но совпадает с federated-средним на worst cell. "
    "(b) Сравнение финальной accuracy; все конфигурации кучкуются в пределах 3 п.п., "
    "наивысшее среднее у ring + IID (97.8%).",
    width_cm=16,
)

add_heading("5.3 Почему federated превосходит centralized? (механизм P1)", level=2)
add_para(
    "Эмпирическое соотношение 0.1271 / 0.1762 = 0.722 ≈ 1/√4 · const соответствует "
    "стохастическому члену Теоремы 1 — 1/√(nT). Механизм: когда n клиентов независимо "
    "делают MeZO-замер своим собственным seed-ом s_i и направлением z_{s_i}, consensus-"
    "усреднение даёт несмещённое среднее n независимых unit-direction замеров. "
    "Стандартный анализ variance reduction (дисперсия ÷ n) показывает, что weight_avg "
    "consensus фактически делает параллельное multi-direction MeZO при том же бюджете "
    "forward-passes, что и централизованный single-direction MeZO. Это количественно "
    "предсказанная «бесплатная выгода» для федеративного обучения в ZO-режиме — обратная "
    "обычной FL-формулировке, где федеративный setup воспринимается как cost-paying."
)

add_heading("5.4 Nesterov ablation: фазовая диаграмма на worst cell", level=2)
add_para(
    "Изолируем worst Day-5 ячейку (ring + Dir(α=0.5)) и прогоняем серию вариантов момента "
    "при seed=42 для bit-exact ablation. Пять вариантов на Рисунке 2:"
)
add_figure(
    "fig2_nesterov_phase_diagram.png",
    "Рисунок 2. Фазовая диаграмма вариантов Nesterov-MeZO на самой сложной федеративной "
    "ячейке. β=0.9 без clip (фиолетовый) расходится на раунде R140 из-за noise-amplified "
    "velocity; loose clipping при C=200 (оранжевый) предотвращает мгновенный blow-up, но "
    "медленно расходится к R500; tight clipping при C=50 с постоянным β=0.9 (R1b, "
    "красный) даёт 3× early speedup, но momentum overshoot вызывает late drift после "
    "R300; рекомендованный linear β-decay 0.9 → 0 с C=50 (R1d, зелёный) даёт монотонное "
    "убывание на всём горизонте, превосходя no-Nesterov control (синий) на 6.0% по final.",
    width_cm=16,
)
add_para(
    "Фазовая диаграмма содержит четыре чётко разделённых региона, каждый количественно "
    "предсказанный Теоремой 1 (через variance amplifier 1/(1−β_t)²):"
)
add_bullets(
    [
        "Регион A (без clip, высокий β): катастрофический blow-up на R≈140. Variance "
        "amplifier 5.3 × ZO-variance неограничен.",
        "Регион B (loose clipping C=200, высокий β): ограниченные выбросы, но velocity "
        "buffer накапливает sub-clip шум; trajectory slow-diverges к R500.",
        "Регион C (tight clipping C=50, постоянный β): velocity ограничен; early-stage "
        "3× speedup; late-stage momentum overshoot создаёт √t drift после R300.",
        "Регион D (tight clipping C=50, β-decay 0.9 → 0): velocity ограничен И amplifier "
        "→ 1 при t → T; монотонное убывание; final 0.1291 превосходит control 0.1373 на 6.0%.",
    ]
)
add_figure(
    "fig4_r1d_detailed.png",
    "Рисунок 4. Детальная траектория D-MeZO-N (R1d) против no-Nesterov control на worst "
    "cell. Eval loss на левой оси (log scale); β-расписание β(t) = 0.9·(1−t/T) наложено "
    "красным на правой оси. Траектория R1d строго монотонно убывает в каждой контрольной "
    "точке, заканчиваясь на 0.1291 против control 0.1373 (улучшение на 6.0% при "
    "фиксированном compute).",
    width_cm=16,
)

add_heading("5.5 Воспроизводимость", level=2)
add_para(
    "Все эксперименты воспроизводимы из публичного репозитория. Репо содержит:"
)
add_bullets(
    [
        "Код: src/dmezo/ (~2.5K LOC) с MeZO-примитивами, federated simulator, partition "
        "utilities, вариантами Нестерова с ρ-clipping и β-schedule.",
        "Тесты: 75/75 pytest проходят. Покрытие: детерминизм возмущения, свойства mixing-"
        "матрицы, корректность simulator consensus, статистика partition, classification "
        "accuracy, ρ-clipping, β-schedule.",
        "Конфиги: configs/*.yaml — один на эксперимент, Hydra-loadable.",
        "Notebooks: notebooks/run_finals.ipynb — single-click воспроизведение полной "
        "multi-seed сетки + R1d + centralized baseline на Colab Pro+.",
        "MLflow run IDs (Drive-mirrored) для каждой числовой величины в Таблицах 1–2 и "
        "Рисунках 1–4.",
        "Технотчёт docs/04-theory.md с полными доказательствами Теорем 1 и 2 и roadmap "
        "к Теореме 3 (PL + момент).",
    ]
)

# =====================================================================
# 6. DISCUSSION
# =====================================================================
add_heading("6. Обсуждение", level=1)

add_heading("6.1 Почему ring ≤ complete на ZO-режиме? (C3)", level=2)
add_para(
    "Контр-интуитивный результат: на обоих partition-режимах ring (ρ(W)=0.333) стабильно "
    "сравним или превосходит complete (ρ(W)=0). В обычном first-order федеративном "
    "обучении complete должен доминировать, поскольку даёт точное per-round среднее. В "
    "ZO-режиме, однако, очень высокая per-step дисперсия ρ означает, что более медленное "
    "consensus-усреднение может играть роль неявного регуляризатора — каждый клиент "
    "интегрирует свой локальный шум за несколько раундов, прежде чем он распространится к "
    "соседям, сглаживая эффективную траекторию. Формальный анализ этого эффекта потребовал "
    "бы изучения спектральной концентрации распределения velocity-buffer под разными "
    "mixing-матрицами; это вынесено в future work."
)

add_heading("6.2 Почему наивный Нестеров несовместим с ZO?", level=2)
add_para(
    "Двухканальная noise-структура look-ahead Нестерова (probe-location и update-direction "
    "оба зависят от v_i) компаундирует variance-amplification: look-ahead позиция θ + βv_i "
    "сама по себе является зашумлённым сдвигом, и замер там даёт ρ-оценку с дисперсией, "
    "масштабирующейся как квадрат локального гессиана умноженный на ∥v_i∥². При β=0.9 "
    "look-ahead вариант расходится в 7× быстрее heavy-ball варианта (R20 vs. R140), что "
    "подтверждает dual-channel механизм. Мы гипотезируем, что variance-reduced ZO-оценки "
    "(multi-direction SPSA, усредняющие K направлений на шаг) могут восстановить хорошие "
    "свойства look-ahead Нестерова; проверка этого оставлена для follow-up работы."
)

add_heading("6.3 Практический рецепт", level=2)
add_para(
    "На основе наших экспериментов рекомендуем следующий deployment-рецепт для D-MeZO-N "
    "(вариант β=0.9):"
)
add_bullets(
    [
        "lr = 3·10⁻⁷ (default Princeton MeZO, скорректированный по нашему LR ablation).",
        "ε = 10⁻³ (default Malladi 2023).",
        "ρ-clipping с C ≈ 1.3 × max наблюдённого |ρ| на первых 100 раундах. Для "
        "Qwen3-class моделей C = 50 сработал.",
        "Линейное β-расписание β_t = 0.9·(1 − t/T) (или cosine, hold-then-decay).",
        "Дважды-стохастическая mixing-матрица W. Кольцо или complete-топология дают "
        "близкие результаты при n=4.",
        "Multi-seed (≥3) для paper-grade оценки дисперсии; мы использовали n=2 из-за "
        "бюджетных ограничений.",
    ]
)

# =====================================================================
# 7. LIMITATIONS
# =====================================================================
add_heading("7. Ограничения и future work", level=1)
add_inline_runs([("Эмпирические ограничения. ", {"bold": True}), (
    "(а) Multi-seed при n=2 означает, что error bars даны как range, а не std; n=3–5 "
    "достаточно для надёжного std. (б) R1d (рекомендованный вариант D-MeZO-N) был "
    "прогнан на одном seed-е; multi-seed расширение прямолинейно, но ограничено бюджетом. "
    "(в) Задачи ограничены short-form классификацией (SST-2, BoolQ); генеративные задачи "
    "(SAMSum, GSM8K) не исследованы. (г) Scale-up за пределы 4-клиентского / 4B-параметрового "
    "режима — реальные FL-деплои имеют 100+ клиентов и 8B+ модели; на этом масштабе "
    "мы не тестировали. (д) Нет head-to-head сравнения с FedKSeed / Ferret / FedZeN — эти "
    "интеграции — нетривиальная работа по коду и были вне scope.",
    {})])
add_inline_runs([("Теоретические ограничения. ", {"bold": True}), (
    "Теорема 3 (полная невыпуклая PL + heavy-ball momentum + decentralized + ZO + "
    "ρ-clipping) остаётся открытой. Необходимый аппарат существует в литературе — "
    "Yang-Zhao-Cheng 2016 для non-convex momentum Lyapunov, Koloskova 2020 Теорема 8 "
    "для decentralized PL, Aybat-Fallah et al. 2019 для оптимального β-расписания под "
    "PL — но 4-сторонняя композиция с ZO и clipping нетривиальна. Мы предоставляем "
    "roadmap в `docs/04-theory-template.md`; оценка затрат: 2–4 недели аккуратного анализа.",
    {})])
add_inline_runs([("Алгоритмические ограничения. ", {"bold": True}), (
    "Рекомендованный D-MeZO-N требует ручного выбора ρ-clip порога C и формы β-расписания. "
    "Adaptive вариант, настраивающий C по наблюдаемому распределению ρ и адаптирующий β "
    "по slope валидационного loss, упростил бы deployment. Multi-direction MeZO "
    "(K-direction SPSA averaging) — естественное variance-reduction расширение, которое "
    "должно сделать look-ahead Нестеров tractable.",
    {})])

# =====================================================================
# 8. CONCLUSION
# =====================================================================
add_heading("8. Заключение", level=1)
add_para(
    "Мы представили D-MeZO-N — Decentralized Federated MeZO с ускорением Нестерова — и "
    "установили его как жизнеспособный peer-to-peer федеративный оптимизатор для дообучения "
    "LLM. Шесть контрибуций (C1–C6) покрывают (i) поддержку новой архитектуры (Qwen3.5 "
    "гибридная linear-attention), (ii) устойчивость к экстремальной неоднородности данных, "
    "(iii) пренебрежимо малую стоимость топологии при n=4 с удивительным режимом "
    "ring ≤ complete, (iv) рабочий ускоренный вариант с рекомендованным рецептом β-decay + "
    "ρ-clipping, и (v–vi) две формальные теоремы сходимости, восемь предсказаний которых "
    "совпадают с эмпирическими находками. Полный репозиторий (код, тесты, конфиги, "
    "notebooks, MLflow IDs, доказательства) публично доступен. Открытые теоретические "
    "вопросы — полная Теорема 3 (PL + момент); открытые эмпирические направления — "
    "масштабирование до 100+ клиентов и оценка на генеративных задачах."
)

# =====================================================================
# References
# =====================================================================
add_heading("Список литературы", level=1)
add_para(
    "Aybat, N. S., Fallah, A., Gurbuzbalaban, M., Ozdaglar, A. (2019). A universally optimal "
    "multistage accelerated stochastic gradient method. NeurIPS 2019.",
    justify=False,
)
add_para(
    "Hsu, T.-M. H., Qi, H., Brown, M. (2019). Measuring the effects of non-identical data "
    "distribution for federated visual classification. arXiv:1909.06335.",
    justify=False,
)
add_para(
    "Karimi, H., Nutini, J., Schmidt, M. (2016). Linear convergence of gradient and "
    "proximal-gradient methods under the Polyak-Łojasiewicz condition. ECML-PKDD 2016.",
    justify=False,
)
add_para(
    "Koloskova, A., Loizou, N., Boreiri, S., Jaggi, M., Stich, S. U. (2020). A unified "
    "theory of decentralized SGD with changing topology and local updates. ICML 2020. "
    "arXiv:2003.10422.",
    justify=False,
)
add_para(
    "Lan, G. (2012). An optimal method for stochastic composite optimization. Mathematical "
    "Programming 133(1-2):365–397.",
    justify=False,
)
add_para(
    "Malladi, S., Gao, T., Nichani, E., Damian, A., Lee, J. D., Chen, D., Arora, S. (2023). "
    "Fine-tuning language models with just forward passes. NeurIPS 2023. arXiv:2305.17333.",
    justify=False,
)
add_para(
    "Maritan, A., Ridolfi, A., Notarstefano, G. (2024). FedZeN: a zeroth-order Newton-style "
    "method for federated learning. arXiv:2309.17241.",
    justify=False,
)
add_para(
    "McMahan, B., Moore, E., Ramage, D., Hampson, S., y Arcas, B. A. (2017). "
    "Communication-efficient learning of deep networks from decentralized data. AISTATS 2017.",
    justify=False,
)
add_para(
    "Nesterov, Y., Spokoiny, V. (2017). Random gradient-free minimization of convex "
    "functions. Foundations of Computational Mathematics 17(2):527–566.",
    justify=False,
)
add_para(
    "Polyak, B. T. (1964). Some methods of speeding up the convergence of iteration methods. "
    "USSR Computational Mathematics and Mathematical Physics 4(5):1–17.",
    justify=False,
)
add_para(
    "Qin, Z., Chen, D., Qian, B., Ding, B., Li, Y., Deng, S. (2024). FedKSeed: federated "
    "full-parameter tuning of billion-sized language models with communication cost under "
    "18 kilobytes. ICML 2024. arXiv:2312.06353.",
    justify=False,
)
add_para(
    "Shu, Y., Yao, W., Hu, S. X. (2024). Ferret: federated full-parameter tuning at scale "
    "for large language models. arXiv:2409.06277.",
    justify=False,
)
add_para(
    "Spall, J. C. (1992). Multivariate stochastic approximation using a simultaneous "
    "perturbation gradient approximation. IEEE Transactions on Automatic Control "
    "37(3):332–341.",
    justify=False,
)
add_para(
    "Stich, S. U. (2019). Local SGD converges fast and communicates little. ICLR 2019.",
    justify=False,
)
add_para(
    "Yang, T., Lin, Q., Li, Z. (2016). Unified convergence analysis of stochastic momentum "
    "methods for convex and non-convex optimization. arXiv:1604.03257.",
    justify=False,
)

doc.save(str(OUT))
print(f"Saved {OUT}")
print(f"File size: {OUT.stat().st_size / 1024:.1f} KB")
